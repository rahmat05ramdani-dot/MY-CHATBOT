import streamlit as st
import random
import google.generativeai as genai
from datetime import datetime, time, timedelta
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from PyPDF2 import PdfReader

# --- 1. Page Configuration and Style ---
st.set_page_config(page_title="Teman Belajar Gemini", page_icon="🎓", layout="wide")

# 🌿 Tema warna hijau lembut
st.markdown(
    """
    <style>
        body, .stApp {
            background-color: #E8F5E9;
        }
        .stTextInput>div>div>input, .stTextArea textarea, .stChatInput textarea {
            background-color: #F1FFF1 !important;
            border-radius: 10px;
            border: 1px solid #A5D6A7;
        }
        .stButton>button {
            background-color: #81C784;
            color: white;
            border-radius: 10px;
            border: none;
        }
        .stButton>button:hover {
            background-color: #66BB6A;
        }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("🎓 Teman Belajar KAMU")
st.caption("Tanya apa saja pasti aku bantu jawab, semangatt belajarnya!")

# --- 2. Sidebar Settings ---
with st.sidebar:
    st.subheader("⚙️ Pengaturan")

    google_api_key = st.text_input("Google AI API Key", type="password")

    st.markdown("### 🎯 Mode Belajar")
    learn_mode = st.radio(
        "Pilih mode belajar:",
        ["Tanya Materi", "Kuis Cepat", "Penjelasan Soal"],
        help="Pilih jenis interaksi belajar yang kamu mau."
    )

    st.markdown("### ⏰ Jadwal Belajar Harian")
    start_time = st.time_input("Mulai Belajar", value=time(19, 0))
    durasi = st.number_input("Durasi Belajar (jam)", min_value=0.5, max_value=6.0, value=2.0, step=0.5)
    end_time = (datetime.combine(datetime.today(), start_time) + timedelta(hours=durasi)).time()

    st.write(
        f"📚 Jadwal kamu: **{start_time.strftime('%H:%M')}** - **{end_time.strftime('%H:%M')}** "
        f"🕓 (Durasi: {durasi} jam)"
    )

    reset_button = st.button("🔄 Reset Percakapan", help="Mulai ulang percakapan dan hapus riwayat belajar.")

# --- 3. API Key Check ---
if not google_api_key:
    st.info("🗝️ Masukkan Google AI API key di sidebar untuk mulai belajar.", icon="💬")
    st.stop()

# Konfigurasi API
genai.configure(api_key=google_api_key)

# --- 4. Analisis PDF ---
st.markdown("## 📄 Analisis File PDF")

uploaded_file = st.file_uploader("Unggah file PDF untuk dianalisis", type=["pdf"])

if uploaded_file is not None:
    pdf_reader = PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""

    st.success(f"✅ PDF berhasil dibaca! Panjang teks: {len(text)} karakter.")
    st.text_area("📜 Cuplikan isi PDF:", text[:1500], height=200)

    if st.button("🔍 Analisis Isi PDF"):
        try:
            system_prompt = (
                "Kamu adalah asisten belajar yang menganalisis dokumen. "
                "Beri ringkasan dan temukan poin-poin penting dari teks berikut:"
            )
            model = genai.GenerativeModel("gemini-2.5-flash")
            response = model.generate_content(f"{system_prompt}\n\n{text}")


            st.subheader("📊 Hasil Analisis")
            st.markdown(response.text)
        except Exception as e:
            st.error(f"Gagal menganalisis PDF: {e}")

# --- 5. Chat Session Setup ---
if "messages" not in st.session_state:
    st.session_state.messages = []
if "history" not in st.session_state:
    st.session_state.history = []

if reset_button:
    st.session_state.clear()
    st.rerun()

# --- 6. Display Past Messages ---
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# --- 7. Chat Input ---
motivasi_list = [
    "Terus semangat belajar ya, Rahmat! 💪",
    "Kegigihanmu hari ini menentukan hasil besok. 📚",
    "Belajar itu proses, bukan perlombaan. 🌱",
    "Satu langkah kecil setiap hari = perubahan besar! 🚀"
]

prompt = st.chat_input("Tulis pertanyaan atau jawab kuis di sini...")

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    try:
        if learn_mode == "Tanya Materi":
            system_prompt = (
                "Kamu adalah asisten belajar yang sabar dan menjelaskan konsep dengan mudah dipahami. "
                "Jawab pertanyaan pengguna dengan jelas, ringkas, dan beri contoh nyata jika bisa."
            )
        elif learn_mode == "Kuis Cepat":
            system_prompt = (
                "Bertindaklah sebagai guru yang memberikan kuis singkat kepada pengguna. "
                "Jika pengguna menjawab, nilai jawaban dan beri umpan balik. "
                "Jika belum ada pertanyaan, berikan satu pertanyaan kuis acak singkat."
            )
        else:
            system_prompt = (
                "Bantu pengguna memahami langkah-langkah menyelesaikan soal secara detail dan logis. "
                "Gunakan format langkah demi langkah agar mudah diikuti."
            )

        model = genai.GenerativeModel("gemini-2.5-flash")
        response = model.generate_content(f"{system_prompt}\n\n{prompt}")
        answer = response.text or "Tidak ada jawaban yang dihasilkan."

        if random.random() < 0.2:
            answer += "\n\n---\n💡 " + random.choice(motivasi_list)

    except Exception as e:
        answer = f"Terjadi kesalahan: {e}"

    with st.chat_message("assistant"):
        st.markdown(answer)

    st.session_state.messages.append({"role": "assistant", "content": answer})
    st.session_state.history.append({
        "mode": learn_mode,
        "question": prompt,
        "answer": answer,
        "timestamp": datetime.now().strftime("%d-%m-%Y %H:%M:%S")
    })

# --- 8. Display History ---
with st.expander("🧾 Lihat Riwayat Belajar"):
    if st.session_state.history:
        for i, item in enumerate(st.session_state.history, 1):
            st.markdown(f"**{i}. Mode:** {item['mode']}")
            st.markdown(f"**Pertanyaan:** {item['question']}")
            st.markdown(f"**Jawaban:** {item['answer']}")
            st.markdown(f"🕒 **Waktu:** {item['timestamp']}")
            st.markdown("---")
    else:
        st.write("Belum ada riwayat belajar.")

# --- 9. Reminder Jadwal ---
now = datetime.now().time()
if start_time <= now <= end_time:
    st.success("⏰ Saatnya belajar! Kamu sedang dalam waktu belajar yang dijadwalkan.")
else:
    st.info("📅 Di luar jam belajar. Jangan lupa istirahat yang cukup, Rahmat 😊")

# --- 10. Export History (PDF / Word) ---
if st.session_state.history:
    st.markdown("### 💾 Simpan Riwayat Belajar")

    col1, col2 = st.columns(2)

    # PDF Export
    with col1:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph("📘 Riwayat Belajar Gemini", styles["Title"]), Spacer(1, 12)]

        for i, item in enumerate(st.session_state.history, 1):
            story.append(Paragraph(f"<b>{i}. Mode:</b> {item['mode']}", styles["Normal"]))
            story.append(Paragraph(f"<b>Pertanyaan:</b> {item['question']}", styles["Normal"]))
            story.append(Paragraph(f"<b>Jawaban:</b> {item['answer']}", styles["Normal"]))
            story.append(Paragraph(f"<b>Waktu:</b> {item['timestamp']}", styles["Normal"]))
            story.append(Spacer(1, 12))

        doc.build(story)
        st.download_button(
            label="⬇️ Unduh sebagai PDF",
            data=buffer.getvalue(),
            file_name="riwayat_belajar.pdf",
            mime="application/pdf"
        )

    # Word Export
    with col2:
        docx_file = BytesIO()
        document = Document()
        document.add_heading("📗 Riwayat Belajar Gemini", level=1)

        for i, item in enumerate(st.session_state.history, 1):
            document.add_paragraph(f"{i}. Mode: {item['mode']}")
            document.add_paragraph(f"Pertanyaan: {item['question']}")
            document.add_paragraph(f"Jawaban: {item['answer']}")
            document.add_paragraph(f"Waktu: {item['timestamp']}")
            document.add_paragraph("-" * 50)

        document.save(docx_file)
        docx_file.seek(0)
        st.download_button(
            label="⬇️ Unduh sebagai Word (DOCX)",
            data=docx_file,
            file_name="riwayat_belajar.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
